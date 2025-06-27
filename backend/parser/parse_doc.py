#!/usr/bin/env python3
"""
Lexi-Fill Document Parser
Parses .docx files and extracts placeholders with smart logic.
"""

import sys
import json
import re
from docx import Document
from pathlib import Path
from collections import defaultdict


def normalize_id(label):
    return re.sub(r'[^a-z0-9]+', '_', label.lower()).strip('_')

def infer_type(raw, label):
    if "date" in label.lower():
        return "date"
    if raw.startswith("$"):
        return "currency"
    if "email" in label.lower():
        return "email"
    if "address" in label.lower():
        return "text"
    return "text"

def extract_inline_label(line, start_pos=0):
    pattern = r'[“"”](.*?)[”"]'
    forward_candidates = []
    backward_candidates = []

    for match in re.finditer(pattern, line):
        label = match.group(1).strip()
        if match.start() >= start_pos:
            dist = match.start() - start_pos
            forward_candidates.append((dist, label))
        else:
            dist = start_pos - match.end()
            backward_candidates.append((dist, label))

    if forward_candidates:
        forward_candidates.sort(key=lambda x: x[0])
        return forward_candidates[0][1]
    elif backward_candidates:
        backward_candidates.sort(key=lambda x: x[0])
        return backward_candidates[0][1]
    
    return None


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    for paragraph in doc.paragraphs:
        text = ''.join(run.text for run in paragraph.runs)
        full_text.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = ''.join(run.text for run in cell.paragraphs[0].runs)
                full_text.append(cell_text)

    full_text = '\n'.join(full_text)
    return full_text

def save_text_to_file(text, original_file_path):
    """Save extracted text to a .txt file for debugging"""
    original_path = Path(original_file_path)
    txt_file_path = original_path.with_suffix('.txt')
    
    try:
        with open(txt_file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Extracted text saved to: {txt_file_path}")
        return str(txt_file_path)
    except Exception as e:
        print(f"Warning: Could not save text file: {e}")
        return None
    
def save_to_file(content, file_path):
        """Save content to a file"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"Error saving file: {e}")
            return False
        

def find_placeholders(text):
    lines = text.splitlines()
    seen_ids = set()
    results = []
    replaced_lines = lines.copy()  

    bracket_patterns = [
        (re.compile(r'\$?\[\s*([^\[\]]{1,}?)\s*\]'), '[]'),                            # $[ _____ ] or $[ \t ]
        # (re.compile(r'\[\s*([^\[\]]*?[^\s\[\]]{2,}[^\[\]]*?)\s*\]'), '[]'),                             # [ PLACEHOLDER ]
        (re.compile(r'\{\{\s*([^{}]*?[^\s{}]{2,}[^{}]*?)\s*\}\}'), '{{}}'),         # {{ PLACEHOLDER }}
        (re.compile(r'\(\(\s*([^()]*?[^\s()]{2,}[^()]*?)\s*\)\)'), '(())'),         # (( PLACEHOLDER ))
        # (re.compile(r'__\s*([^_]*?[^\s_]{2,}[^_]*?)\s*__'), '__ __'),               # __ PLACEHOLDER __
    ]


    for idx, line in enumerate(lines):
        current_line = replaced_lines[idx] 
        
        line_matches = []
        for pattern, raw_fmt in bracket_patterns:
            matches = list(pattern.finditer(current_line))
            for match_obj in matches:
                line_matches.append((match_obj, raw_fmt))
        
        line_matches.sort(key=lambda x: x[0].start())
        
        for match_obj, raw_fmt in line_matches:
            match_text = match_obj.group(1) if match_obj.groups() else match_obj.group(0)
            placeholder_start = match_obj.start()
            
            is_generic = not match_text.strip() or re.fullmatch(r'[_\s\t]+', match_text)
            inline_label = extract_inline_label(line, placeholder_start) if is_generic else None
            label = inline_label or match_text.strip()

            raw = raw_fmt
            ph_id = normalize_id(label)

            if ph_id not in seen_ids:
                seen_ids.add(ph_id)
                
                results.append({
                    "raw": raw,
                    "label": label,
                    "id": ph_id,
                    "type": infer_type(raw, label),
                    "line_num": idx
                })
        
        for match_obj, raw_fmt in reversed(line_matches):
            match_text = match_obj.group(1) if match_obj.groups() else match_obj.group(0)
            placeholder_start = match_obj.start()
            
            is_generic = not match_text.strip() or re.fullmatch(r'[_\s\t]+', match_text)
            inline_label = extract_inline_label(line, placeholder_start) if is_generic else None
            label = inline_label or match_text.strip()
            ph_id = normalize_id(label)
            
            if raw_fmt == '[]' and match_obj.group(0)[0] == '$':
                replacement_token = f" $<<{ph_id}>>"
                replaced_lines[idx] = (
                    current_line[:match_obj.start()-1] 
                    + replacement_token
                    + current_line[match_obj.end():]
                )
            else:
                replacement_token = f"<<{ph_id}>>"
                replaced_lines[idx] = (
                    current_line[:match_obj.start()]
                    + replacement_token
                    + current_line[match_obj.end():]
                )
            
            current_line = replaced_lines[idx]

    label_colon_pattern = re.compile(r'^\s*([A-Za-z][\w\s&\-]{0,40}[a-z][\w\s&\-]*):(?=\s*$|\t)')
    section_heading_pattern = re.compile(r'^\s*([A-Z][A-Z\s&\-]{1,40})[:\s]*$')
    current_section = None

    for idx, line in enumerate(lines):
        heading_match = section_heading_pattern.match(line)
        if heading_match:
            current_section = normalize_id(heading_match.group(1).strip())
        
        for label_match in re.finditer(label_colon_pattern, line):    
            label = label_match.group(1).strip()
            base_id = normalize_id(label)
            scoped_id = f"{current_section}_{base_id}" if current_section else base_id

            if scoped_id not in seen_ids:
                seen_ids.add(scoped_id)
                results.append({
                "raw": label_match.group(1),
                "label": label,
                "id": scoped_id,
                "type": infer_type("____", label),
                # "placeholder_num": placeholder_ct,
                "line_num": idx
                })
                # placeholder_ct += 1
            replacement_token = f"<<{scoped_id}>>"
            colon_index = label_match.end()
            replaced_lines[idx] = (
                line[:colon_index] + replacement_token + line[colon_index:]
            )
                

    return results, '\n'.join(replaced_lines)

def main():
    if len(sys.argv) != 2:
        print(json.dumps({"error": "Usage: python3 parse_doc.py <file_path>"}))
        sys.exit(1)

    file_path = sys.argv[1]
    if not Path(file_path).exists():
        print(json.dumps({"error": f"[ERROR] File not found: {file_path}"}))
        sys.exit(1)

    try:
        text = extract_text_from_docx(file_path)        
        placeholders, replaced_template = find_placeholders(text)
        print(json.dumps({
            "placeholders": placeholders,
            "replacedTemplate": replaced_template
            }, indent=2))

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
